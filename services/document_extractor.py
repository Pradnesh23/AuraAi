"""
Document extraction service
Handles ZIP extraction, PDF conversion, DOCX parsing, and OCR text extraction
"""
import zipfile
import tempfile
import shutil
import re
import uuid
from pathlib import Path
from typing import List, Dict, Optional, Tuple
import logging
from concurrent.futures import ThreadPoolExecutor, as_completed

import cv2
import numpy as np
import pytesseract
from PIL import Image
from pdf2image import convert_from_path

try:
    from docx import Document as DocxDocument
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False

from config import UPLOAD_DIR, PROCESSED_DIR, ALLOWED_EXTENSIONS, OCR_LANGUAGES
from services.image_processor import ImageProcessor

logger = logging.getLogger(__name__)


class DocumentExtractor:
    """Handles document extraction and OCR processing"""
    
    def __init__(self):
        self.image_processor = ImageProcessor()
        self.session_id = str(uuid.uuid4())
    
    def save_individual_file(self, content: bytes, filename: str) -> Tuple[str, Path]:
        """
        Save an individual uploaded file
        
        Args:
            content: File content as bytes
            filename: Original filename
            
        Returns:
            Tuple of (session_id, saved file path)
        """
        session_dir = UPLOAD_DIR / self.session_id
        session_dir.mkdir(parents=True, exist_ok=True)
        
        safe_name = self._sanitize_filename(filename)
        target_path = session_dir / safe_name
        
        with open(target_path, 'wb') as f:
            f.write(content)
        
        logger.info(f"Saved individual file: {safe_name}")
        return self.session_id, target_path
    
    def save_multiple_files(self, files: List[Tuple[bytes, str]]) -> Tuple[str, List[Path]]:
        """
        Save multiple individual files
        
        Args:
            files: List of (content, filename) tuples
            
        Returns:
            Tuple of (session_id, list of saved file paths)
        """
        session_dir = UPLOAD_DIR / self.session_id
        session_dir.mkdir(parents=True, exist_ok=True)
        
        saved_paths = []
        for content, filename in files:
            ext = Path(filename).suffix.lower()
            if ext not in ALLOWED_EXTENSIONS:
                logger.warning(f"Skipping unsupported file: {filename}")
                continue
            
            safe_name = self._sanitize_filename(filename)
            target_path = session_dir / safe_name
            
            with open(target_path, 'wb') as f:
                f.write(content)
            
            saved_paths.append(target_path)
            logger.info(f"Saved: {safe_name}")
        
        return self.session_id, saved_paths
    
    def extract_zip(self, zip_path: Path) -> Tuple[str, List[Path]]:
        """
        Extract ZIP file and return list of valid document paths
        
        Args:
            zip_path: Path to the uploaded ZIP file
            
        Returns:
            Tuple of (session_id, list of extracted file paths)
        """
        session_dir = UPLOAD_DIR / self.session_id
        session_dir.mkdir(parents=True, exist_ok=True)
        
        extracted_files = []
        
        try:
            with zipfile.ZipFile(zip_path, 'r') as zip_ref:
                for file_info in zip_ref.infolist():
                    # Skip directories
                    if file_info.is_dir():
                        continue
                    
                    # Skip hidden files and __MACOSX
                    filename = Path(file_info.filename).name
                    if filename.startswith('.') or '__MACOSX' in file_info.filename:
                        continue
                    
                    # Check file extension
                    ext = Path(filename).suffix.lower()
                    if ext not in ALLOWED_EXTENSIONS:
                        logger.warning(f"Skipping unsupported file: {filename}")
                        continue
                    
                    # Secure extraction - prevent path traversal
                    safe_name = self._sanitize_filename(filename)
                    target_path = session_dir / safe_name
                    
                    # Extract file
                    with zip_ref.open(file_info) as source:
                        with open(target_path, 'wb') as target:
                            shutil.copyfileobj(source, target)
                    
                    extracted_files.append(target_path)
                    logger.info(f"Extracted: {safe_name}")
            
            return self.session_id, extracted_files
            
        except zipfile.BadZipFile as e:
            logger.error(f"Invalid ZIP file: {e}")
            raise ValueError("Invalid or corrupted ZIP file")
    
    def _sanitize_filename(self, filename: str) -> str:
        """Sanitize filename to prevent security issues"""
        # Remove path components
        name = Path(filename).name
        # Remove potentially dangerous characters
        name = re.sub(r'[^\w\-_\.]', '_', name)
        # Ensure unique name
        return f"{uuid.uuid4().hex[:8]}_{name}"
    
    def process_document(self, file_path: Path) -> Dict:
        """
        Process a single document and extract text
        
        Args:
            file_path: Path to the document (PDF, DOCX, or image)
            
        Returns:
            Dictionary with extracted data
        """
        ext = file_path.suffix.lower()
        
        if ext == '.pdf':
            return self._process_pdf(file_path)
        elif ext in ['.docx', '.doc']:
            return self._process_docx(file_path)
        else:
            return self._process_image(file_path)
    
    def _process_docx(self, docx_path: Path) -> Dict:
        """Extract text from DOCX file"""
        try:
            if not DOCX_SUPPORT:
                raise ValueError("DOCX support not available. Install python-docx.")
            
            doc = DocxDocument(str(docx_path))
            
            # Extract all paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text.strip())
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        paragraphs.append(" | ".join(row_text))
            
            full_text = "\n".join(paragraphs)
            candidate_name = self._extract_candidate_name(full_text, docx_path.stem)
            
            return {
                "id": str(uuid.uuid4()),
                "name": candidate_name,
                "source_file": docx_path.name,
                "raw_text": full_text,
                "page_count": 1
            }
            
        except Exception as e:
            logger.error(f"Error processing DOCX {docx_path}: {e}")
            raise
    
    def _process_pdf(self, pdf_path: Path) -> Dict:
        """Convert PDF to images and extract text"""
        try:
            # Poppler path for Windows
            poppler_path = r"C:\Users\prads\Downloads\Release-25.12.0-0\poppler-25.12.0\Library\bin"
            
            # Convert PDF pages to images (reduced DPI for speed)
            images = convert_from_path(
                pdf_path,
                dpi=150,  # Reduced from 300 for faster processing
                fmt='png',
                poppler_path=poppler_path
            )
            
            all_text = []
            
            for i, pil_image in enumerate(images):
                # Convert PIL Image to numpy array
                np_image = np.array(pil_image)
                # Convert RGB to BGR for OpenCV
                if len(np_image.shape) == 3:
                    np_image = cv2.cvtColor(np_image, cv2.COLOR_RGB2BGR)
                
                # Preprocess the image
                processed = self.image_processor.preprocess(np_image)
                
                # Extract text
                text = self._ocr_image(processed)
                all_text.append(text)
                
                logger.debug(f"Processed page {i+1} of {pdf_path.name}")
            
            full_text = "\n\n".join(all_text)
            candidate_name = self._extract_candidate_name(full_text, pdf_path.stem)
            
            return {
                "id": str(uuid.uuid4()),
                "name": candidate_name,
                "source_file": pdf_path.name,
                "raw_text": full_text,
                "page_count": len(images)
            }
            
        except Exception as e:
            logger.error(f"Error processing PDF {pdf_path}: {e}")
            raise
    
    def _process_image(self, image_path: Path) -> Dict:
        """Process a single image file"""
        try:
            # Load image
            image = self.image_processor.load_image(image_path)
            if image is None:
                raise ValueError(f"Could not load image: {image_path}")
            
            # Preprocess
            processed = self.image_processor.preprocess(image)
            
            # OCR
            text = self._ocr_image(processed)
            candidate_name = self._extract_candidate_name(text, image_path.stem)
            
            return {
                "id": str(uuid.uuid4()),
                "name": candidate_name,
                "source_file": image_path.name,
                "raw_text": text,
                "page_count": 1
            }
            
        except Exception as e:
            logger.error(f"Error processing image {image_path}: {e}")
            raise
    
    def _ocr_image(self, image: np.ndarray) -> str:
        """
        Perform OCR on preprocessed image using Tesseract
        """
        # Configure Tesseract path for Windows
        pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
        
        # Configure Tesseract for best results
        custom_config = r'--oem 3 --psm 6'
        
        try:
            text = pytesseract.image_to_string(
                image,
                lang='+'.join(OCR_LANGUAGES),
                config=custom_config
            )
            return text.strip()
        except Exception as e:
            logger.error(f"OCR failed: {e}")
            return ""
    
    def _extract_candidate_name(self, text: str, fallback: str) -> str:
        """
        Attempt to extract candidate name from resume text
        Falls back to filename if extraction fails
        """
        if not text:
            return self._clean_fallback_name(fallback)
        
        lines = text.split('\n')
        
        # Common patterns for names at the start of resumes
        for line in lines[:10]:  # Check first 10 lines
            line = line.strip()
            if not line or len(line) < 3:
                continue
            
            # Skip lines that look like headers or contact info
            skip_patterns = [
                r'^(resume|cv|curriculum)',
                r'^(phone|email|address|linkedin)',
                r'@',  # Email
                r'\d{5,}',  # Phone numbers
                r'^(objective|summary|experience)',
            ]
            
            should_skip = False
            for pattern in skip_patterns:
                if re.search(pattern, line.lower()):
                    should_skip = True
                    break
            
            if should_skip:
                continue
            
            # Check if line looks like a name (2-4 words, mostly letters)
            words = line.split()
            if 1 <= len(words) <= 4:
                # Check if mostly alphabetic
                alpha_ratio = sum(c.isalpha() or c.isspace() for c in line) / len(line)
                if alpha_ratio > 0.8:
                    return line.title()
        
        return self._clean_fallback_name(fallback)
    
    def _clean_fallback_name(self, filename: str) -> str:
        """Clean up filename to use as candidate name"""
        # Remove UUID prefix if present
        name = re.sub(r'^[a-f0-9]{8}_', '', filename)
        # Remove common resume suffixes
        name = re.sub(r'[-_]?(resume|cv|final|v\d+)$', '', name, flags=re.IGNORECASE)
        # Replace separators with spaces
        name = re.sub(r'[-_]', ' ', name)
        return name.title()
    
    def process_all_documents(
        self, 
        file_paths: List[Path],
        max_workers: int = 4
    ) -> List[Dict]:
        """
        Process all extracted documents in parallel
        
        Args:
            file_paths: List of paths to process
            max_workers: Maximum number of parallel workers
            
        Returns:
            List of extracted document data
        """
        results = []
        
        # Use ThreadPoolExecutor for parallel processing
        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            # Submit all tasks
            future_to_path = {
                executor.submit(self.process_document, path): path 
                for path in file_paths
            }
            
            # Collect results as they complete
            for future in as_completed(future_to_path):
                path = future_to_path[future]
                try:
                    result = future.result()
                    results.append(result)
                    logger.info(f"Successfully processed: {path.name}")
                except Exception as e:
                    logger.error(f"Failed to process {path.name}: {e}")
                    continue
        
        return results
    
    def cleanup_session(self, session_id: str):
        """Remove temporary files for a session"""
        session_dir = UPLOAD_DIR / session_id
        if session_dir.exists():
            shutil.rmtree(session_dir)
            logger.info(f"Cleaned up session: {session_id}")
